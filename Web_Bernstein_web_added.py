# Web_20250312.py
import streamlit as st
from datetime import datetime
from datetime import timedelta
import re
import io
import DifyNews_API as DifyNews
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from pathlib import Path
import io
# from streamlit_autorefresh import st_autorefresh #自动刷新——不用这个功能

# os.getcwd()
# os.chdir(r'e:\coding\vscode\html编程')
# # FONT_DIR = Path(__file__).parent / "fonts"
# path=os.path.join(os.getcwd(),'fonts')
# # pdfmetrics.registerFont(TTFont('SimSun', str(FONT_DIR / 'SimSun.ttf')))
# # pdfmetrics.registerFont(TTFont('SimHei', str(FONT_DIR / 'SimHei.ttf')))
# pdfmetrics.registerFont(TTFont('SimSun', path + '/SimSun.ttf'))
# pdfmetrics.registerFont(TTFont('SimHei', path + '/SimHei.ttf'))
pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttf'))
pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))

# # 注册中文字体（需确保系统有对应字体文件）
# def init_fonts():
#     try:
#         pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
#         pdfmetrics.registerFont(TTFont('SimHei', 'SimHei.ttf'))
#     except:
#         st.warning("中文字体加载失败，请确保系统已安装SimSun字体")


def generate_pdf(content, filename):
    pdf_io = io.BytesIO()
    c = canvas.Canvas(pdf_io, pagesize=letter)
  
    # 设置中文字体和编码
    c.setFont("SimSun", 12)
  
    # 处理多行文本
    y_position = 750
    for line in content.split('\n'):
        # 自动换行处理
        text = c.beginText(50, y_position)
        text.setFont("SimSun", 12)
        text.textLine(line.strip())
        c.drawText(text)
        y_position -= 20
      
        if y_position < 50:
            c.showPage()
            y_position = 750
  
    c.save()
    pdf_io.seek(0)
    return pdf_io


# ----------------------
# 金融风格设计系统
# ----------------------
FINANCE_THEME = {
    "primary": "#1F4172",    # 主蓝
    "secondary": "#132043",  # 深蓝
    "accent": "#F1B4BB",     # 强调色
    "neutral": "#F0F3FF"     # 背景灰
}

def apply_finance_style():
    """注入金融行业标准样式"""
    st.markdown(f"""
    <style>
    .main {{
        background-color: {FINANCE_THEME['neutral']};
    }}
    .stButton>button {{
        border-radius: 8px!important;
        background: {FINANCE_THEME['primary']}!important;
        color: white!important;
        border: 1px solid {FINANCE_THEME['secondary']}!important;
    }}
    .report-title {{
        color: {FINANCE_THEME['secondary']};
        border-left: 4px solid {FINANCE_THEME['accent']};
        padding-left: 1rem;
    }}
    </style>
    """, unsafe_allow_html=True)

# ----------------------
# 核心业务模块
# ----------------------
def init_session():
    """初始化会话状态"""
    session_defaults = {
        'page': "主页面",
        'generated_content': "",
        'keywords': []
    }
    for key, val in session_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

def generate_document(file_type, content, file_name):
    """通用文档生成器"""
    if file_type == "Word":
        from docx import Document
        doc = Document()
        doc.add_heading(file_name, level=1)
        doc.add_paragraph(content)
      
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
  
    elif file_type == "PDF":
        from reportlab.pdfgen import canvas
        pdf_io = io.BytesIO()
        c = canvas.Canvas(pdf_io, pagesize=(600, 800))
        c.setFont("Helvetica", 12)
      
        # 专业PDF排版
        y_position = 750
        for line in content.split('\n'):
            c.drawString(50, y_position, line)
            y_position -= 20
            if y_position < 50:
                c.showPage()
                y_position = 750
        c.save()
        pdf_io.seek(0)
        return pdf_io, "application/pdf"

# ----------------------
# 新闻生成页面
# ----------------------
def render_news_generator():
    with st.container():
        st.markdown('<h2 class="report-title">金融快讯生成系统</h2>', unsafe_allow_html=True)
      
        # 参数输入区
        col1, col2 = st.columns([1, 3])
        input_params = {"tone": None,
                        "word_count": None,
                        "event": None,
                        "reference": None,
                        "original_content": None,
                        "restatement_objective": None,  
                        "language": None,
                        "article_type": None}
        with col1:
            function = st.selectbox('功能模式', ['转述', '总结', '生成'], 
                                 format_func=lambda x: f"📌 {x}")
            if function == "生成":
                input_params['article_type'] = st.selectbox("文档类型", ['news', 'feature', 'commentary'])
            
            input_params['tone']=st.selectbox("语气", ['formal', 'informal', 'persuasive', 'neutral'])
            input_params['word_count'] = st.slider("目标字数", 100, 2000, 100)
        # 动态输入配置
        
        with col2:
            if function == "转述":
                input_params['original_content'] = st.text_area("原文输入", height=150)
                input_params['restatement_objective'] = st.text_input("改写要求")
              
            elif function == "总结":
                input_params['original_content'] = st.text_area("长文本输入", height=200)
              
            elif function == "生成":
                input_params['event'] = st.text_area("事件描述", 
                    # placeholder="输入金融事件关键要素：\n- 涉及机构\n- 金额规模\n- 政策影响", 
                    height=150)
                input_params['reference'] = st.text_area("参考网站" )
            input_params['language'] = st.selectbox("输出语言", ['Chinese', 'English'])

         
        if "history" not in st.session_state:
            st.session_state["history"] = []
        result=None
        # 生成控制
        if st.button("立即生成", use_container_width=True):
            try:
                need="Article Restatement" if function == "转述" else "Article summary" if function == "总结" else 'Article Generation' if function == "生成" else "Briefing"
                result = DifyNews.Run_Dify({
                    "need": need,
                    **input_params
                })
                print(result)
                st.session_state["output_type"] = 'Restated Article' if function == "转述" else 'Article Summary' if function == "总结" else "Generated Article" if function == "生成" else 'weekly_report'
                st.session_state["news"]=result.get('data', {}).get('outputs', {}).get(st.session_state["output_type"], '')
                
            #历史记录
                # st.session_state.history.append({"tries": "{0}_{1}".format(st.session_state["output_type"],datetime.now().strftime('%Y%m%d_%H%M%S')), "content" : result.get('data', {}).get('outputs', {}).get(st.session_state["output_type"], '') })
            
                
                if result.get("error"):
                    # st.session_state.history.append({"role": "User", "content": "failed to generate"})
                    st.error(f"生成失败：{result['error']}")

                else:
                    # timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    # st.session_state.history.append({"type": function,
                    #     "content": result.get('data', {}).get('outputs', {}).get(st.session_state["output_type"], ''),
                    #     "timestamp": timestamp
                    #     })
                    # ----------------------
                    # 在新闻生成成功处修改历史记录添加方式
                    # timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    # st.session_state.history.append({
                    #     "type": function,  # 记录类型：转述/总结/生成
                    #     "content": result.get('data', {}).get('outputs', {}).get(st.session_state["output_type"], ''),
                    #     "timestamp": timestamp,
                    #     "metadata": {
                    #         "word_count": input_params['word_count'],
                    #         "language": input_params['language'],
                    #         "tone": input_params.get('tone'),
                    #         "article_type": input_params.get('article_type')
                    #     }
                    # })
                    # print(st.session_state.history)
                    st.success(f"生成成功：{st.session_state['output_type']}")
            except Exception as e:
                st.error(f"系统错误：{str(e)}")

        # 结果展示与导出
        if 'news' in st.session_state:
            with st.expander("生成结果", expanded=True):
                st.markdown("""
                            <style>
                                .stExpander pre {
                                     white-space: pre-wrap !important;
                                     word-wrap: break-word !important;
                                    }  
                            </style>
                            """, unsafe_allow_html=True)
                st.markdown(f"```\n{st.session_state["news"]}\n```")
          
            # export_type = st.radio("导出格式", ["PDF", "Word"], horizontal=True)
            file_name = st.text_input("文件名/标题（可选）", value=f"{st.session_state["output_type"]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}", key='file_name_input')
            export_format = st.selectbox("选择导出格式", ["Word", "PDF"], key='export_format_input')
            col1, col2 = st.columns(2)
            with col1:
                if st.button("生成可下载文件"):
                    if export_format == "Word":
                        doc = Document()
                        doc.add_heading(file_name, level=1)
                        doc.add_paragraph(st.session_state["news"]) 

                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)  
                        st.download_button("下载 Word 文件", doc_io, file_name + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    elif export_format == "PDF":
                        pdf_io = generate_pdf(st.session_state["news"], file_name)
                        st.download_button(
                        "下载 PDF 文件",
                        data=pdf_io,
                        file_name=f"{file_name}.pdf",
                        mime="application/pdf"
                        )
            with col2:
                if st.button("保存到历史记录"):
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    st.session_state.history.append({
                        "type": function,  # 记录类型：转述/总结/生成
                        "content": st.session_state["news"],
                        "timestamp": timestamp,
                        "metadata": {
                            "word_count": input_params['word_count'],
                            "language": input_params['language'],
                            "tone": input_params.get('tone'),
                            "article_type": input_params.get('article_type')
                        }
                    })
                    # count = st_autorefresh(interval=20, limit=1, key="fizzbuzzcounter")
                    # print(st.session_state.history)
                    st.success("历史记录保存成功！")

# 在原有代码基础上新增/修改以下部分

# ----------------------
# 周报生成页面
# ----------------------
def render_weekly_report():
    with st.container():
        st.markdown('<h2 class="report-title">金融监管周报系统</h2>', unsafe_allow_html=True)
        result=None
        # 时间范围选择
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("起始日期", 
                                      value=datetime.now() - timedelta(days=7),
                                      min_value=datetime(2023,1,1))
        with col2:
            end_date = st.date_input("结束日期", 
                                    value=datetime.now(),
                                    max_value=datetime.now())
        language = st.selectbox("输出语言", ['Chinese', 'English'])
        # 关键词输入
        keywords = st.text_area("监管关键词（逗号分隔）", 
                              placeholder="输入关键词示例：\n- 风险管理\n- 资本充足率\n- 反洗钱",
                              height=100)
        keyword_list = [kw.strip() for kw in re.split(r'[，,、]', keywords) if kw.strip()]
      
        # 生成参数设置
        with st.expander("高级设置"):
            article_count = st.slider("最大文章数", 5, 50, 10)
            risk_level = st.select_slider("风险等级", 
                                         options=["常规", "关注", "高度关注", "紧急"])
        if "history" not in st.session_state:
            st.session_state["history"] = []
        # 生成控制
        if st.button("生成周报", use_container_width=True):
            try:
                # 参数校验
                if not keyword_list:
                    st.warning("请至少输入一个关键词")
                    return
              
                # 构造请求参数
                # input_text = {
                #     "need": "Briefing",
                #     "start_date": start_date.strftime("%Y-%m-%d"),
                #     "end_date": end_date.strftime("%Y-%m-%d"),
                #     "keywords": keyword_list,
                #     "risk_level": risk_level,
                #     "article_count": article_count,
                #     "language": "Chinese"
                # }
                input_text = {
                    "need": "Briefing",  
                    "tone": None, 
                    "word_count": 200, 
                    "event":', '.join(keyword_list), 
                    "reference": None,
                    "original_content": None,
                    "restatement_objective": None,  
                    "language": language,  
                    "article_type": None
                    }
              
                # 调用API
                with st.spinner("⏳ 正在生成专业周报..."):
                    result = DifyNews.Run_Dify(input_text)
                    print(result)
                # 处理结果
                if "error" in result:
                    
                    st.error(f"生成失败：{result['error']}")
                else:
                    st.session_state['weekly_report'] = result.get('data', {}).get('outputs', {}).get('weekly_report', '')
                    st.success("周报生成成功！")
                  
                    # 记录生成历史
                    # timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    # st.session_state.history.append({
                    #     "type": "周报",
                    #     "date_range": f"{start_date}至{end_date}",
                    #     "content": st.session_state['result'],
                    #     "timestamp": timestamp
                    # })
                    

                    # 在周报生成成功处修改历史记录添加方式
                    
            except Exception as e:
                st.error(f"系统错误：{str(e)}")

        # 结果展示与导出
        if 'weekly_report' in st.session_state and st.session_state['weekly_report']:
            with st.expander("周报预览", expanded=True):
                st.markdown(f"```\n{st.session_state['weekly_report']}\n```")
          
            # 导出选项
            export_type = st.radio("导出格式", ["PDF", "Word"], horizontal=True, key='weekly_export')
            file_name = st.text_input("周报文件名", 
                                    value=f"监管周报_{datetime.now().strftime('%Y%m%d')}", 
                                    key='weekly_filename')
            col1, col2 = st.columns(2)
            with col2:
                if st.button("保存到历史记录"):
                    st.session_state.history.append({
                        "type": "周报",
                        "content": st.session_state['weekly_report'],
                        "timestamp": datetime.now().strftime('%Y%m%d_%H%M%S'),
                        "date_range": f"{start_date}至{end_date}",
                        "metadata": {
                            "keywords": keyword_list,
                            "risk_level": risk_level,
                            "article_count": article_count
                        }
                    })
                    # count = st_autorefresh(interval=20, limit=1, key="fizzbuzzcounter")
                    st.success("历史记录保存成功！")
            with col1:
                if st.button("生成可下载文件"):
                    if export_type == "Word":
                        doc = Document()
                        doc.add_heading(file_name, 0)
                        doc.add_paragraph(st.session_state.weekly_report)
                  
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                  
                        st.download_button(
                            label="下载Word周报",
                            data=doc_io,
                            file_name=f"{file_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        pdf_io = generate_pdf(st.session_state.weekly_report, file_name)
                        st.download_button(
                            "下载PDF周报",
                            data=pdf_io,
                            file_name=f"{file_name}.pdf",
                            mime="application/pdf"
                        )
# if 'history' in st.session_state:
#     print(st.session_state.history)
# ------------------------
# ----------------------
# 在侧边栏添加历史记录模块
# ----------------------
def show_history_panel():
    """侧边栏历史记录面板"""
    with st.sidebar:
        st.divider()
        st.markdown("**历史记录**")
      
        if not st.session_state.get('history'):
            st.info("暂无生成记录")
            return

        # 生成历史记录选项标签
        st.session_state['history_options'] = [
            f"{h['type']}_{h['timestamp']}" 
            for h in st.session_state.history 
            if 'type' in h and 'timestamp' in h
        ]
        print(st.session_state['history_options'])
        selected_record = st.selectbox(
            "选择历史记录",
            options=st.session_state['history_options'],
            format_func=lambda x: x.replace('_', ' '),
            key='history_selector'
        )
      
        if selected_record:
            # 获取选中记录的完整数据
            record_index = st.session_state['history_options'].index(selected_record)
            print(record_index,'\n')
            selected_data = st.session_state.history[record_index]
            print(selected_data,'\n')
          
            # 显示元数据
            with st.expander("记录详情", expanded=True):
                cols = st.columns(2)
                cols[0].caption(f"类型：{selected_data['type']}")
                cols[1].caption(f"生成时间：{selected_data['timestamp']}")
                if 'date_range' in selected_data:
                    st.caption(f"时间范围：{selected_data['date_range']}")
          
            # 导出控制
            export_type = st.radio("导出格式", 
                                 ["Word", "PDF"],
                                 horizontal=True,
                                 key='history_export_type')
          
            if st.button("导出选中记录"):
                file_name = f"{selected_data['type']}_{selected_data['timestamp']}"
                content = selected_data['content']
              
                if export_type == "Word":
                    doc = Document()
                    doc.add_heading(file_name, level=1)
                    doc.add_paragraph(content)
                  
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                  
                    st.download_button(
                        label="下载Word文件",
                        data=doc_io,
                        file_name=f"{file_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success("文件导出成功！")
                else:
                    pdf_io = generate_pdf(content, file_name)
                    st.download_button(
                        label="下载PDF文件",
                        data=pdf_io,
                        file_name=f"{file_name}.pdf",
                        mime="application/pdf",
                    )
                    st.success("文件导出成功！")    
                    

# ----------------------
# 修改生成函数以记录完整元数据


# ----------------------
# 在主程序入口添加历史面板调用
# ----------------------
if __name__ == "__main__":
    init_session()
    apply_finance_style()

    # 侧边栏布局
    with st.sidebar:
        nav_choice = st.radio("导航菜单", ["新闻生成", "监管周报"])
    
    # 主内容区保持不变
    if nav_choice == "新闻生成":
        render_news_generator()
    else:
        render_weekly_report()

    with st.sidebar:
        # 新增历史记录面板
        show_history_panel()
  
# 在 Web_Bernstein_web_added.py 中插入
# ... existing code ...
# ----------------------
# 在页面底部插入图片框
st.markdown(""" 
    <div style="text-align: left; margin-top: 20px;">
        <h6 style="margin: 0;">📖我们的数据来源（点击快捷访问）</h6>
        <div style="display: flex; justify-content: space-between; margin-top: 10px;">
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://www.nfra.gov.cn/cn/view/pages/index/index.html" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://img.bjd.com.cn/p/2023/05/18/adb532c7973ef7ff240ca678f55b3c58.png" alt="中国国家金融监督管理总局" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">国家金融监管总局</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://finance.sina.com.cn/money/insurance/" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://pic4.zhimg.com/v2-8565f43d2bc3740409e64c50266fdfe7_r.jpg" alt="b" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">新浪财经保险频道</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="http://www.cbimc.cn/" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://51huoke.oss-cn-shenzhen.aliyuncs.com/image/2020_6/2/766ddbf1c06df7e1a141f66734fb49d7jpeg" alt="c" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">中国银行保险报网</p>
            </div>
        </div>
        <div style="display: flex; justify-content: space-between; margin-top: 10px;">
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://www.iachina.cn/" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://n.sinaimg.cn/sinakd20201119s/217/w1080h737/20201119/3cbe-kcysmrw9716527.png" alt="d" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">中国保险行业协会</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="http://finance.ce.cn/insurance/" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://th.bing.com/th/id/R.26c88f5e30df955787adabc702aeb672?rik=IOHD3inp8ep3cg&riu=http%3a%2f%2fi.ce.cn%2fvr%2fcy%2f201610%2f08%2fW020161101772223254001.jpg&ehk=5W4rOBsDY0PMMFg7jYOazrMcRoFLmu1cNLVI6FDLPIA%3d&risl=&pid=ImgRaw&r=0" alt="e" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">中国经济网保险频道</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://insurance.eastmoney.com/" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://th.bing.com/th/id/R.ce1dacad1a0b9edf2069d3c35c44b081?rik=N%2b4jxEX85Bqe6A&riu=http%3a%2f%2fimg.ji7.net%2fuploads%2f5dc01a81176d4.jpg&ehk=IrxXbLkpPCh4x3G%2fK8LpgX59XQ6wdZ8EeWd6Y%2fuJX6s%3d&risl=&pid=ImgRaw&r=0" alt="e" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">东方财富网保险头条</p>
            </div>
        </div>
        <div style="display: flex; justify-content: space-between; margin-top: 10px;">
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://www.financialnews.com.cn/node_3007.html" target="_blank">
                    <div style="width: 100%; height: 120px; overflow: hidden; border-radius: 10px;">
                        <img src="https://img.xieniao.com/jietu/20231117/wwwfinancialnewscomcn.jpg" alt="f" style="width: 100%; height: auto; object-fit: cover;">
                    </div>
                </a>
                <p style="margin: 5px 0;">中国金融新闻网保险频道</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; margin-right: 10px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://www.eeo.com.cn/eeo/caijing/baoxian/" target="_blank">
                    <img src="https://pic.rmb.bdstatic.com/ca27d2c3fe7f79833542f69dce028de7.png@wm_2,t_55m+5a625Y+3L+eQvOmfs+S8oOWqkg==,fc_ffffff,ff_U2ltSGVp,sz_18,x_12,y_12" alt="g" style="width: 100%; height: 120px; object-fit: cover;">
                </a>
                <p style="margin: 5px 0;">经济观察网保险频道</p>
            </div>
            <div style="border: 2px solid lightgray; border-radius: 10px; background-color: white; padding: 2px; text-align: center; flex: 1; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <a href="https://business.sohu.com/" target="_blank">
                    <img src="https://5b0988e595225.cdn.sohucs.com/images/20180817/11fe3784a0a24b1a878126375a972c91.jpeg" alt="h" style="width: 100%; height: 120px; object-fit: cover;">
                </a>
                <p style="margin: 5px 0;">搜狐财经</p>
            </div>
        </div>
    </div>
""", unsafe_allow_html=True)
# ... existing code ...


# ----------------------

